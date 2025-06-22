import os
import json
from pathlib import Path
import PyPDF2
from pypdf import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import time
import docx
from docx import Document as DocxDocument

class PDFProcessor:
    def __init__(self, pdf_folder_path="pdfs"):
        self.pdf_folder_path = Path(pdf_folder_path)
        self.processed_documents = []
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text from a single PDF file"""
        try:
            print(f"Processing PDF: {pdf_path}")
            
            # Try with pypdf first (better for most PDFs)
            try:
                reader = PdfReader(str(pdf_path))
                text = ""
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text
                    except Exception as e:
                        print(f"Error extracting page {page_num + 1} from {pdf_path}: {e}")
                        continue
                
                if text.strip():
                    return text
                    
            except Exception as e:
                print(f"pypdf failed for {pdf_path}, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    
                    for page_num in range(len(pdf_reader.pages)):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                text += f"\n--- Page {page_num + 1} ---\n"
                                text += page_text
                        except Exception as e:
                            print(f"Error extracting page {page_num + 1} from {pdf_path}: {e}")
                            continue
                    
                    return text
                    
            except Exception as e:
                print(f"PyPDF2 also failed for {pdf_path}: {e}")
                return None
                
        except Exception as e:
            print(f"Error processing PDF {pdf_path}: {e}")
            return None
    
    def extract_text_from_word(self, word_path):
        """Extract text from a Word document (.docx)"""
        try:
            print(f"Processing Word document: {word_path}")
            
            # Load the Word document
            doc = DocxDocument(str(word_path))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text if text.strip() else None
            
        except Exception as e:
            print(f"Error processing Word document {word_path}: {e}")
            return None
    
    def clean_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Skip lines that are just page numbers or headers/footers
                if len(line) > 3 and not line.isdigit():
                    cleaned_lines.append(line)
        
        # Join lines and normalize spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove multiple consecutive newlines
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        
        return cleaned_text
    
    def process_single_document(self, file_path):
        """Process a single document (PDF or Word)"""
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        # Determine file type and extract text accordingly
        if file_extension == '.pdf':
            raw_text = self.extract_text_from_pdf(file_path)
            source_type = 'pdf'
        elif file_extension in ['.docx', '.doc']:
            if file_extension == '.doc':
                print(f"⚠️  .doc format not supported for {file_path.name}. Please convert to .docx")
                return []
            raw_text = self.extract_text_from_word(file_path)
            source_type = 'word'
        else:
            print(f"⚠️  Unsupported file format: {file_extension} for {file_path.name}")
            return []
        
        if not raw_text or len(raw_text.strip()) < 100:
            print(f"⚠️  Insufficient text extracted from {file_path.name}")
            return []
        
        # Clean the text
        cleaned_text = self.clean_text(raw_text)
        
        if not cleaned_text:
            print(f"⚠️  No valid text after cleaning from {file_path.name}")
            return []
        
        # Create a document object
        doc = Document(
            page_content=cleaned_text,
            metadata={
                'source': str(file_path),
                'filename': file_path.name,
                'source_type': source_type,
                'timestamp': time.time(),
                'character_count': len(cleaned_text)
            }
        )
        
        # Split the document into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Add chunk metadata
        for j, chunk in enumerate(chunks):
            chunk.metadata.update({
                'chunk_id': f"{file_path.stem}_chunk_{j}",
                'chunk_index': j,
                'total_chunks': len(chunks)
            })
        
        print(f"✓ Created {len(chunks)} chunks from {file_path.name}")
        return chunks
    
    def process_all_pdfs(self):
        """Process all PDFs and Word documents in the folder"""
        if not self.pdf_folder_path.exists():
            print(f"Document folder {self.pdf_folder_path} does not exist!")
            return []
        
        # Find all supported document files
        pdf_files = list(self.pdf_folder_path.glob("*.pdf"))
        word_files = list(self.pdf_folder_path.glob("*.docx"))
        doc_files = list(self.pdf_folder_path.glob("*.doc"))
        
        all_files = pdf_files + word_files + doc_files
        
        if not all_files:
            print(f"No supported document files found in {self.pdf_folder_path}")
            print("Supported formats: .pdf, .docx")
            return []
        
        print(f"Found {len(all_files)} document files to process:")
        print(f"  - PDFs: {len(pdf_files)}")
        print(f"  - Word documents (.docx): {len(word_files)}")
        if doc_files:
            print(f"  - Word documents (.doc): {len(doc_files)} (will be skipped - please convert to .docx)")
        
        all_documents = []
        
        for i, file_path in enumerate(all_files):
            print(f"\nProcessing {i+1}/{len(all_files)}: {file_path.name}")
            
            # Process the document
            chunks = self.process_single_document(file_path)
            all_documents.extend(chunks)
        
        self.processed_documents = all_documents
        print(f"\n=== DOCUMENT PROCESSING SUMMARY ===")
        print(f"Total files processed: {len([f for f in all_files if f.suffix.lower() != '.doc'])}")
        print(f"Total chunks created: {len(all_documents)}")
        print(f"Average chunk size: {sum(len(doc.page_content) for doc in all_documents) // len(all_documents) if all_documents else 0} characters")
        
        # Show breakdown by file type
        pdf_chunks = [doc for doc in all_documents if doc.metadata['source_type'] == 'pdf']
        word_chunks = [doc for doc in all_documents if doc.metadata['source_type'] == 'word']
        
        if pdf_chunks:
            print(f"PDF chunks: {len(pdf_chunks)}")
        if word_chunks:
            print(f"Word document chunks: {len(word_chunks)}")
        
        return all_documents
    
    def save_processed_data(self, filename="processed_documents.json"):
        """Save processed document data to JSON file"""
        try:
            # Convert documents to serializable format
            serializable_docs = []
            for doc in self.processed_documents:
                serializable_docs.append({
                    'content': doc.page_content,
                    'metadata': doc.metadata
                })
            
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(serializable_docs, f, ensure_ascii=False, indent=2)
            
            print(f"Processed document data saved to {filename}")
            
        except Exception as e:
            print(f"Error saving processed data: {e}")
    
    def load_processed_data(self, filename="processed_documents.json"):
        """Load processed document data from JSON file"""
        try:
            if not os.path.exists(filename):
                print(f"File {filename} does not exist")
                return []
            
            with open(filename, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert back to Document objects
            documents = []
            for item in data:
                doc = Document(
                    page_content=item['content'],
                    metadata=item['metadata']
                )
                documents.append(doc)
            
            self.processed_documents = documents
            print(f"Loaded {len(documents)} processed documents from {filename}")
            return documents
            
        except Exception as e:
            print(f"Error loading processed data: {e}")
            return []

def main():
    """Main function to test PDF and Word document processing"""
    # Initialize processor
    processor = PDFProcessor("pdfs")  # Assumes documents are in 'pdfs' folder
    
    try:
        # Process all documents
        documents = processor.process_all_pdfs()
        
        if documents:
            # Save processed data
            processor.save_processed_data("processed_documents.json")
            
            # Show sample chunks from different document types
            print(f"\nSample chunks:")
            
            # Show PDF sample if available
            pdf_docs = [doc for doc in documents if doc.metadata['source_type'] == 'pdf']
            if pdf_docs:
                print(f"\nPDF Sample:")
                sample_pdf = pdf_docs[0]
                print(f"Source: {sample_pdf.metadata['filename']}")
                print(f"Chunk ID: {sample_pdf.metadata['chunk_id']}")
                print(f"Content preview: {sample_pdf.page_content[:200]}...")
            
            # Show Word document sample if available
            word_docs = [doc for doc in documents if doc.metadata['source_type'] == 'word']
            if word_docs:
                print(f"\nWord Document Sample:")
                sample_word = word_docs[0]
                print(f"Source: {sample_word.metadata['filename']}")
                print(f"Chunk ID: {sample_word.metadata['chunk_id']}")
                print(f"Content preview: {sample_word.page_content[:200]}...")
        else:
            print("No documents were processed successfully.")
            
    except Exception as e:
        print(f"Error during document processing: {e}")

if __name__ == "__main__":
    main()